import os
import base64
import re
import time
from pathlib import Path
from typing import List, Dict, Any, Optional

from src.utils import find_data_directory
import config
from src.utils.logging_config import setup_logger
from src.utils.report_formatter import (
    sanitize_report_for_display,
    sanitize_user_visible_text,
    format_investigation_result
)
from src.utils.report_generator import generate_report_llm

logger = setup_logger(__name__)

def validate_image_path(image_path: str) -> tuple[bool, Optional[str], Optional[str]]:
    """이미지 파일 경로 유효성 검사

    Args:
        image_path: 검증할 이미지 파일 경로

    Returns:
        (is_valid, error_msg, resolved_path) 튜플
        - 유효하지 않은 경우: (False, error_msg, None)
        - 유효하며 원본 경로에 존재: (True, None, None)
        - 유효하며 data_dir에만 존재: (True, None, str(alt_path))
    """
    image_file = Path(image_path)

    # 파일 존재 확인
    if not image_file.exists():
        # data 폴더에서도 확인
        try:
            data_dir = find_data_directory()
            alt_path = Path(data_dir) / image_file.name
            if alt_path.exists():
                return True, None, str(alt_path)
        except (OSError, ValueError):
            pass
        return False, f"이미지 파일을 찾을 수 없습니다: {image_path}", None

    # 파일 크기 검증
    try:
        MAX_IMAGE_SIZE = config.MAX_IMAGE_SIZE_MB * 1024 * 1024
        file_size = image_file.stat().st_size
        if file_size > MAX_IMAGE_SIZE:
            return False, (
                f"이미지 파일이 너무 큽니다 ({file_size / 1024 / 1024:.1f}MB). "
                f"최대 {config.MAX_IMAGE_SIZE_MB}MB까지 지원합니다."
            ), None
    except OSError as e:
        return False, f"파일 크기 확인 실패: {e}", None

    # 파일 확장자 검증 (config.IMAGE_EXTENSIONS와 동기화, 대소문자 무관)
    # 예: "*.png" -> lstrip('*')=".png", "*.PNG" -> ".png" (lower)
    ext = image_file.suffix.lower()
    valid_extensions = list({
        e.lstrip('*').lower() for e in config.IMAGE_EXTENSIONS
    })
    if ext not in valid_extensions:
        return False, f"지원하지 않는 이미지 형식입니다: {ext}", None

    return True, None, None

def validate_result_structure(result: Any) -> tuple[bool, Optional[str]]:
    """분석 결과 딕셔너리 구조 검증
    
    Args:
        result: 검증할 결과 객체
        
    Returns:
        (유효성 여부, 오류 메시지) 튜플
        - 유효한 경우: (True, None)
        - 유효하지 않은 경우: (False, 오류 메시지)
    """
    if result is None:
        return False, "분석 결과를 받지 못했습니다."
    
    if not isinstance(result, dict):
        return False, f"예상치 못한 결과 타입: {type(result)}"
    
    # 필수 키 확인
    required_keys = ["final_verdict", "expert_reports", "arbiter_debate_messages", "errors"]
    for key in required_keys:
        if key not in result:
            return False, f"필수 키가 누락되었습니다: {key}"
    
    # 타입 검증
    if not isinstance(result.get("expert_reports"), list):
        return False, "expert_reports가 리스트가 아닙니다."
    
    if not isinstance(result.get("arbiter_debate_messages"), list):
        return False, "arbiter_debate_messages가 리스트가 아닙니다."
    
    if not isinstance(result.get("errors"), list):
        return False, "errors가 리스트가 아닙니다."
    
    return True, None

def select_image_file() -> Optional[str]:
    """data 디렉토리에서 이미지 파일 목록을 보여주고 사용자가 선택할 수 있게 합니다.
    
    Returns:
        선택된 이미지 파일 경로 또는 None (취소/오류 시)
    """
    try:
        data_dir = find_data_directory()
    except ValueError as e:
        logger.error(f"데이터 디렉토리 찾기 실패: {e}")
        print(f"오류: {e}")
        return None
    except Exception as e:
        logger.exception(f"예상치 못한 오류 발생: {e}")
        print(f"오류: {e}")
        return None
    
    try:
        image_extensions = config.IMAGE_EXTENSIONS
        image_files = []
        for ext in image_extensions:
            try:
                image_files.extend(Path(data_dir).glob(ext))
            except OSError as e:
                logger.warning(f"이미지 파일 검색 중 오류 (확장자: {ext}): {e}")
                continue
        
        image_files = sorted(set(image_files))
        
        if not image_files:
            logger.warning(f"이미지 파일 없음: {data_dir}")
            print(f"오류: {data_dir} 디렉토리에 이미지 파일이 없습니다.")
            return None
        
        print("\n" + "=" * 60)
        print("사용 가능한 이미지 파일:")
        print("=" * 60)
        for idx, img_file in enumerate(image_files, 1):
            try:
                file_size = img_file.stat().st_size / 1024  # KB
                print(f"  [{idx}] {img_file.name} ({file_size:.1f} KB)")
            except OSError as e:
                logger.warning(f"파일 정보 읽기 실패: {img_file.name}, 오류: {e}")
                print(f"  [{idx}] {img_file.name} (크기 확인 불가)")
        print("=" * 60)
        
        while True:
            try:
                choice = input(f"\n이미지 번호를 선택하세요 (1-{len(image_files)}): ").strip()
                
                if choice.lower() in ['q', 'quit']:
                    logger.info("사용자가 선택 취소")
                    print("실행을 취소했습니다.")
                    return None
                
                choice_num = int(choice)
                if 1 <= choice_num <= len(image_files):
                    selected_file = image_files[choice_num - 1]
                    logger.info(f"이미지 선택됨: {selected_file.name}")
                    print(f"\n선택된 이미지: {selected_file.name}")
                    return str(selected_file)
                else:
                    print(f"오류: 1부터 {len(image_files)} 사이의 숫자를 입력해주세요.")
            except ValueError:
                print("오류: 숫자를 입력해주세요. (종료하려면 'q' 입력)")
            except KeyboardInterrupt:
                logger.info("사용자가 키보드 인터럽트로 취소")
                print("\n\n실행을 취소했습니다.")
                return None
            except Exception as e:
                logger.exception(f"이미지 선택 중 예상치 못한 오류: {e}")
                print(f"오류: 예상치 못한 문제가 발생했습니다. 다시 시도해주세요.")
    except Exception as e:
        logger.exception(f"이미지 파일 목록 생성 중 오류: {e}")
        print(f"오류: 이미지 파일 목록을 생성하는 중 오류가 발생했습니다: {e}")
        return None

def create_payload_from_image(image_path: str) -> List[Dict[str, Any]]:
    """
    이미지 경로에서 직접 Gemini Vertex AI 형식의 payload 생성
    """
    # 이미지 파일 유효성 검사
    is_valid, error_msg, resolved_path = validate_image_path(image_path)
    if not is_valid:
        if "찾을 수 없습니다" in error_msg:
            raise FileNotFoundError(error_msg)
        else:
            raise ValueError(error_msg)

    path_to_use = resolved_path or image_path
    image_file = Path(path_to_use)

    # 이미지 데이터 읽기
    try:
        with open(path_to_use, 'rb') as f:
            image_data = f.read()
    except PermissionError as exc:
        logger.error(f"파일 읽기 권한 없음: {path_to_use}")
        raise PermissionError(f"이미지 파일 읽기 권한이 없습니다: {path_to_use}") from exc
    except OSError as e:
        logger.error(f"파일 읽기 오류: {path_to_use}, 오류: {e}")
        raise ValueError(f"이미지 파일 읽기 중 오류 발생: {e}") from e
    except Exception as e:
        logger.exception(f"예상치 못한 오류 발생: {path_to_use}, 오류: {e}")
        raise ValueError(f"이미지 파일 처리 중 예상치 못한 오류 발생: {e}") from e

    # MIME 타입 결정
    ext = image_file.suffix.lower()
    if ext == '.png':
        mime_type = 'image/png'
    elif ext in ['.jpg', '.jpeg']:
        mime_type = 'image/jpeg'
    elif ext in ['.heic']:
        mime_type = 'image/heic'
    else:
        # 기본값으로 jpeg 사용
        mime_type = 'image/jpeg'
    
    try:
        image_base64 = base64.b64encode(image_data).decode('utf-8')
    except Exception as e:
        logger.error(f"base64 인코딩 실패: {image_path}, 오류: {e}")
        raise ValueError(f"이미지 base64 인코딩 중 오류 발생: {e}") from e
    
    payload = [
        {
            "text": "이미지를 분석하고 화재 원인을 조사하세요. (분석 단계별 지침을 따르세요)"
        },
        {
            "inline_data": {
                "mime_type": mime_type,
                "data": image_base64
            }
        }
    ]
    return payload

def _get_expert_filename(report: str) -> str:
    """전문가 리포트에서 파일명 추출
    
    지원하는 리포트 형식:
    - [Contact 전문가 최종 판정 - Analyst-Critic 토론] (base_debate_nodes)
    - ## Contact 전문가 최종 분석 결과 (verdict_debate_nodes)
    """
    # 패턴 1: [Expert 전문가 ...] 형식
    if "[Contact 전문가" in report or "## Contact 전문가" in report:
        return "Contact_Expert_Report.txt"
    elif "[Deform 전문가" in report or "## Deform 전문가" in report:
        return "Deform_Expert_Report.txt"
    elif "[Necking 전문가" in report or "## Necking 전문가" in report:
        return "Necking_Expert_Report.txt"
    elif "[Aging 전문가" in report or "## Aging 전문가" in report:
        return "Aging_Expert_Report.txt"
    elif "[DielectricAge 전문가" in report or "## DielectricAge 전문가" in report:
        return "Dielectric_Expert_Report.txt"
    elif "[Mechanical 전문가" in report or "## Mechanical 전문가" in report:
        return "Mechanical_Expert_Report.txt"
    elif "[StrandFracture 전문가" in report or "## StrandFracture 전문가" in report:
        return "StrandFracture_Expert_Report.txt"
    elif "[Tracking 전문가" in report or "## Tracking 전문가" in report:
        return "Tracking_Expert_Report.txt"
    return "Unknown_Expert_Report.txt"

def save_expert_reports(expert_reports: List[str], output_dir: Path) -> None:
    """각 전문가별 별도 리포트 파일 저장
    동일 전문가가 여러 개일 경우 Contact_Expert_Report.txt, Contact_Expert_Report_2.txt 등으로 구분.
    """
    if not expert_reports:
        return

    logger.info("개별 리포트 저장 시작")
    expert_count: Dict[str, int] = {}
    for report in expert_reports:
        base_name = _get_expert_filename(report)
        expert_count[base_name] = expert_count.get(base_name, 0) + 1
        if expert_count[base_name] > 1:
            stem = base_name.replace(".txt", "")
            filename = f"{stem}_{expert_count[base_name]}.txt"
        else:
            filename = base_name
        expert_file = output_dir / filename
        try:
            expert_file.parent.mkdir(parents=True, exist_ok=True)
            with open(expert_file, 'w', encoding='utf-8') as f:
                f.write(sanitize_report_for_display(report))
            logger.debug(f"전문가 리포트 저장 완료: {filename}")
            print(f"  - {filename} 저장 완료")
        except (OSError, UnicodeEncodeError) as e:
            logger.error(f"전문가 리포트 저장 실패: {filename}, 오류: {e}")
            print(f"  - {filename} 저장 실패: {e}")

def save_arbiter_report(final_verdict: str, output_dir: Path) -> None:
    """Arbiter(최종 결론) 리포트 별도 저장"""
    if not final_verdict:
        return
    
    arbiter_file = output_dir / "Arbiter_Report.txt"
    try:
        arbiter_file.parent.mkdir(parents=True, exist_ok=True)
        with open(arbiter_file, 'w', encoding='utf-8') as f:
            f.write("[Arbiter (Chief Investigator) Report]\n\n")
            f.write(final_verdict)
        logger.debug("Arbiter 리포트 저장 완료")
        print("  - Arbiter_Report.txt 저장 완료")
    except (OSError, UnicodeEncodeError) as e:
        logger.error(f"Arbiter 리포트 저장 실패: {e}")
        print(f"  - Arbiter_Report.txt 저장 실패: {e}")

def _to_relative_image_path(output_file: Path, visual_report_path: str) -> str:
    """output_file 기준으로 visual_report_path의 상대 경로 반환 (마크다운 이미지용)"""
    try:
        out_dir = Path(output_file).parent.resolve()
        img_path = Path(visual_report_path).resolve()
        rel = img_path.relative_to(out_dir)
        return str(rel).replace('\\', '/')
    except (ValueError, OSError):
        # 경로 계산 실패 시 파일명만 사용 (같은 outputs 하위 구조 가정)
        return f"../visual_reports/{Path(visual_report_path).name}"


def _to_embedded_image_data_uri(visual_report_path: str) -> Optional[str]:
    """이미지 파일을 Base64로 읽어 data URI 반환 (마크다운에 직접 삽입용)"""
    try:
        img_path = Path(visual_report_path)
        if not img_path.exists():
            return None
        with open(img_path, 'rb') as f:
            raw = f.read()
        b64 = base64.b64encode(raw).decode('ascii')
        ext = img_path.suffix.lower()
        mime = 'image/jpeg' if ext in ('.jpg', '.jpeg') else 'image/png' if ext == '.png' else 'image/jpeg'
        return f"data:{mime};base64,{b64}"
    except (OSError, ValueError) as e:
        logger.warning(f"이미지 Base64 인코딩 실패: {e}")
        return None


def _resolve_image_for_docx(
    img_ref: str,
    base_dir: Path,
    raw_visual_report_path: Optional[str] = None,
) -> Optional[Path]:
    """이미지 참조를 실제 파일 경로로 해석 (docx 삽입용)"""
    if not img_ref or not img_ref.strip():
        return None
    img_ref = img_ref.strip()
    # data: URI는 별도 처리
    if img_ref.startswith("data:"):
        return None
    # file:/// 경로
    if img_ref.lower().startswith("file:///"):
        p = Path(img_ref[8:].lstrip("/"))
        if p.exists():
            return p
        # Windows: file:///C:/... -> C:\...
        try:
            p_win = Path(img_ref.replace("file:///", "").replace("/", "\\"))
            if p_win.exists():
                return p_win
        except Exception:
            pass
        return None
    # 상대 경로 (../visual_reports/xxx.jpg 등)
    resolved = (base_dir / img_ref.replace("\\", "/")).resolve()
    if resolved.exists():
        return resolved
    # raw_visual_report_path 직접 사용
    if raw_visual_report_path:
        rp = Path(raw_visual_report_path)
        if rp.exists():
            return rp
        # cwd 기준
        cwd_path = Path.cwd() / raw_visual_report_path.replace("\\", "/")
        if cwd_path.exists():
            return cwd_path
    return None


def _setup_docx_styles(doc) -> None:
    """Word 문서 기본 스타일 설정 (가독성 개선)"""
    from docx.shared import Pt

    font_name = "Malgun Gothic"

    def _set_style_font(font, size_pt: int):
        font.size = Pt(size_pt)
        font.name = font_name

    # 본문 (Normal): 11pt
    _set_style_font(doc.styles["Normal"].font, 11)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.15  # 1.15배

    # 제목 1 (H1): 18pt
    if "Title" in doc.styles:
        _set_style_font(doc.styles["Title"].font, 18)
    for h in ["Heading 1", "heading 1"]:
        if h in doc.styles:
            _set_style_font(doc.styles[h].font, 18)
            doc.styles[h].paragraph_format.space_before = Pt(12)
            doc.styles[h].paragraph_format.space_after = Pt(6)
            break

    # 제목 2 (H2): 14pt
    for h in ["Heading 2", "heading 2"]:
        if h in doc.styles:
            _set_style_font(doc.styles[h].font, 14)
            doc.styles[h].paragraph_format.space_before = Pt(10)
            doc.styles[h].paragraph_format.space_after = Pt(4)
            break

    # 제목 3 (H3): 12pt
    for h in ["Heading 3", "heading 3"]:
        if h in doc.styles:
            _set_style_font(doc.styles[h].font, 12)
            doc.styles[h].paragraph_format.space_before = Pt(6)
            doc.styles[h].paragraph_format.space_after = Pt(2)
            break


def save_investigation_result_docx(
    formatted_result: str,
    output_docx_path: Path,
    base_dir: Path,
    raw_visual_report_path: Optional[str] = None,
) -> None:
    """포맷된 분석 결과를 Word(.docx) 문서로 저장 (이미지 포함)"""
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    _setup_docx_styles(doc)
    # 페이지 여백 설정
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    # 마크다운 스타일 파싱: ##, ###, 본문, 이미지
    img_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
    lines = formatted_result.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # # 제목 (H1)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:].strip()
            if text:
                doc.add_heading(text, level=0)
            i += 1
            continue
        # ## 섹션 (H2)
        if stripped.startswith("## ") and not stripped.startswith("### "):
            text = stripped[3:].strip()
            if text:
                doc.add_heading(text, level=1)
            i += 1
            continue
        # ### 서브섹션 (H3)
        if stripped.startswith("### "):
            text = stripped[4:].strip()
            if text:
                doc.add_heading(text, level=2)
            i += 1
            continue
        # 테이블 행 (| ... |)
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i])
                i += 1
            if table_lines:
                _add_markdown_table_to_doc(doc, table_lines)
            continue
        # 이미지 단독 행 ![alt](path)
        # 대용량 base64(수MB) 줄은 regex 생략 → 메모리/성능 이슈 방지
        path = None
        if len(line) < 500_000:
            img_match = img_pattern.search(line)
            if img_match and (len(line) < 1000 or line.strip() == line):
                path = img_match.group(2)
        else:
            if "![" in line and "](" in line and ")" in line and "base64," in line:
                start = line.find("](") + 2
                end = line.rfind(")")
                if end > start:
                    path = line[start:end]
        if path:
            if path.startswith("data:"):
                # Base64: regex 대신 문자열 슬라이싱 (수MB 문자열에서 메모리/성능 이슈 방지)
                try:
                    import io
                    idx = path.find("base64,")
                    if idx >= 0:
                        b64_data = path[idx + 7:]
                        raw = base64.b64decode(b64_data)
                        doc.add_picture(io.BytesIO(raw), width=Inches(5.0))
                except Exception as e:
                    logger.warning(f"Base64 이미지 docx 삽입 실패: {e}")
            else:
                img_path = _resolve_image_for_docx(path, base_dir, raw_visual_report_path)
                if img_path:
                    try:
                        doc.add_picture(str(img_path), width=Inches(5.0))
                    except Exception as e:
                        logger.warning(f"이미지 docx 삽입 실패 {img_path}: {e}")
            i += 1
            continue
        # 구분선 ---
        if stripped in ("---", "***", "___"):
            i += 1
            continue
        # 일반 단락 (수MB base64 줄은 위에서 처리되므로 여기선 제외)
        if stripped and len(stripped) < 100_000:
            # 인라인 이미지가 있는 경우 분리
            parts = img_pattern.split(line)
            if len(parts) > 1:
                for j, part in enumerate(parts):
                    if j % 3 == 0 and part.strip():
                        doc.add_paragraph(part.strip())
                    elif j % 3 == 2:
                        img_path = _resolve_image_for_docx(part, base_dir, raw_visual_report_path)
                        if img_path:
                            try:
                                doc.add_picture(str(img_path), width=Inches(5.0))
                            except Exception:
                                pass
            else:
                doc.add_paragraph(stripped)
        i += 1

    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))


def _add_markdown_table_to_doc(doc, table_lines: List[str]) -> None:
    """마크다운 테이블을 docx에 추가 (가독성 스타일 적용)"""
    from docx.shared import Pt

    if len(table_lines) < 2:
        return
    rows = []
    for ln in table_lines:
        cells = [c.strip() for c in ln.split("|")[1:-1]]
        if not cells:
            continue
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for ri, row_cells in enumerate(rows):
        for ci, cell_text in enumerate(row_cells):
            if ci < col_count:
                cell = table.rows[ri].cells[ci]
                cell.text = cell_text.replace("<br>", "\n").replace("<br/>", "\n")
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after = Pt(2)
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Malgun Gothic"
    doc.add_paragraph()


def save_investigation_result(
    final_verdict: str,
    expert_reports: List[str],
    arbiter_debate_messages: List[Dict[str, Any]],
    errors: List[str],
    input_image_path: str,
    output_file: Path,
    final_verdict_structured: Optional[Any] = None,
    visual_report_path: Optional[str] = None
) -> None:
    """통합 분석 결과 파일 저장"""
    timestamp_str = time.strftime('%Y-%m-%d %H:%M:%S')
    
    # 리포트 생성: LLM 기반(통합 출력만) 또는 정규식 기반
    llm_report = None
    if config.USE_LLM_REPORT_GENERATOR:
        llm_report = generate_report_llm(
            final_verdict=final_verdict,
            expert_reports=expert_reports or [],
            arbiter_debate_messages=arbiter_debate_messages or [],
            input_image_path=input_image_path,
            timestamp=timestamp_str,
            errors=errors,
            final_verdict_structured=final_verdict_structured,
        )
        if llm_report:
            if visual_report_path:
                if getattr(config, 'EMBED_IMAGE_IN_MARKDOWN', False):
                    img_ref = _to_embedded_image_data_uri(visual_report_path) or _to_relative_image_path(output_file, visual_report_path)
                else:
                    img_ref = _to_relative_image_path(output_file, visual_report_path)
                llm_report += f"\n\n## 3. 상세 증거 분석 시각화\n![분석 결과 시각화]({img_ref})\n"
            formatted_result = llm_report
            logger.info("LLM 기반 통합 리포트 생성 완료")
            print("  [리포트] LLM 기반 통합 리포트 생성 완료")
        else:
            if visual_report_path:
                img_ref = _to_embedded_image_data_uri(visual_report_path) if getattr(config, 'EMBED_IMAGE_IN_MARKDOWN', False) else _to_relative_image_path(output_file, visual_report_path)
            else:
                img_ref = None
            formatted_result = format_investigation_result(
                final_verdict=final_verdict,
                expert_reports=expert_reports or [],
                arbiter_debate_messages=arbiter_debate_messages or [],
                input_image_path=input_image_path,
                timestamp=timestamp_str,
                final_verdict_structured=final_verdict_structured,
                visual_report_path=img_ref,
            )
            logger.warning("LLM 리포트 생성 실패, 구조화된 데이터 또는 정규식 Fallback 사용")
            print("  [리포트] LLM 실패 → 구조화된 데이터 또는 정규식 Fallback 사용")
    else:
        if visual_report_path:
            img_ref = _to_embedded_image_data_uri(visual_report_path) if getattr(config, 'EMBED_IMAGE_IN_MARKDOWN', False) else _to_relative_image_path(output_file, visual_report_path)
        else:
            img_ref = None
        formatted_result = format_investigation_result(
            final_verdict=final_verdict,
            expert_reports=expert_reports or [],
            arbiter_debate_messages=arbiter_debate_messages or [],
            input_image_path=input_image_path,
            timestamp=timestamp_str,
            final_verdict_structured=final_verdict_structured,
            visual_report_path=img_ref,
        )

    # LLM 실패 시 또는 비-LLM 모드일 때만 errors 수동 추가
    if errors and not llm_report:
        formatted_result += "\n\n---\n\n## System Errors & Warnings\n\n"
        for err in errors:
            formatted_result += f"- {sanitize_user_visible_text(err)}\n"

    # .txt 저장
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(formatted_result)

    # .docx 저장 (Word 문서)
    docx_path = output_file.with_suffix('.docx')
    try:
        save_investigation_result_docx(
            formatted_result=formatted_result,
            output_docx_path=docx_path,
            base_dir=output_file.parent,
            raw_visual_report_path=visual_report_path,
        )
        logger.debug(f"Word 문서 저장 완료: {docx_path}")
        print(f"  - {docx_path.name} 저장 완료")
    except Exception as e:
        logger.warning(f"Word 문서 저장 실패: {e}", exc_info=True)
        print(f"  ⚠ Word 문서 저장 실패: {e}")
